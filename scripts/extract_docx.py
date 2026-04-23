#!/usr/bin/env python3
"""
Extract text sections from .docx files in assets/docs and produce a JSON file
at src/data/content_from_docs.json

Usage:
  - Place your .docx files in assets/docs/
  - Run: python scripts/extract_docx.py

This script uses python-docx (pip install python-docx).
It detects headings by paragraph style (Heading 1/2/3 or Title) and splits the
document into sections. If no headings are found it creates a single section
with the full text.
"""
import os
import json
import re
from pathlib import Path

try:
    from docx import Document
except Exception as e:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    raise


def slugify(s: str) -> str:
    s = s.lower()
    s = re.sub(r"[^a-z0-9]+", "-", s)
    s = s.strip("-")
    return s or "doc"


def extract_sections(doc_path: Path):
    doc = Document(doc_path)
    sections = []
    current_title = None
    current_lines = []

    def push_section():
        nonlocal current_title, current_lines
        if current_title is None and not current_lines:
            return
        title = current_title or "Resumen"
        text = "\n\n".join([l for l in current_lines if l.strip()])
        sections.append({"title": title, "text": text})
        current_title = None
        current_lines = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            # preserve paragraph breaks
            continue

        style_name = ""
        try:
            style_name = getattr(p.style, "name", "") or ""
        except Exception:
            style_name = ""

        # treat Heading styles or Title as section separators
        if re.match(r"^heading\s*\d+", style_name, re.I) or style_name.lower() in ("title", "subtitle"):
            # push previous
            push_section()
            current_title = text
            current_lines = []
        else:
            current_lines.append(text)

    # push last
    push_section()

    # if no sections found, make a single one with full text
    if not sections:
        full = "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
        sections = [{"title": "Contenido", "text": full}]

    return sections


def main():
    base = Path(__file__).resolve().parents[1]
    docs_dir = base / "assets" / "docs"
    out_file = base / "src" / "data" / "content_from_docs.json"

    if not docs_dir.exists():
        print(f"No se encontró la carpeta {docs_dir}. Crea assets/docs y coloca allí los .docx.")
        return

    topics = []
    files = sorted([p for p in docs_dir.iterdir() if p.suffix.lower() == ".docx"])
    if not files:
        print(f"No hay archivos .docx en {docs_dir}")
        return

    for f in files:
        print(f"Procesando: {f.name}")
        sections = extract_sections(f)
        title = sections[0]["title"] if sections else f.stem
        topic = {
            "id": slugify(f.stem),
            "title": title,
            "sourceFile": str(f.name),
            "sections": sections
        }
        topics.append(topic)

    payload = {"importedFromDocx": True, "topics": topics}
    out_file.parent.mkdir(parents=True, exist_ok=True)
    with open(out_file, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, indent=2, ensure_ascii=False)

    print(f"Exportado {len(topics)} documentos a {out_file}")


if __name__ == "__main__":
    main()
